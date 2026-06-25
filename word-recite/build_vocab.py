import json
import re
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT.parent / "zoology-review" / "zoology_source.docx"


def main():
    doc = Document(SOURCE)
    chapter_lines = [
        p.text.strip() for p in doc.paragraphs
        if re.match(r"^Chapter\s+\d+\s+", p.text.strip(), re.I)
    ]
    chapters = []
    all_words = []
    for index, table in enumerate(doc.tables[1:19], start=1):
        line = chapter_lines[index - 1]
        match = re.match(r"^Chapter\s+(\d+)\s+(.+)$", line, re.I)
        number, title = match.groups()
        words = []
        for row in table.rows[1:]:
            term, meaning, note = [cell.text.strip() for cell in row.cells[:3]]
            if not term:
                continue
            word = {
                "id": f"zoo-{index}-{len(words)+1}",
                "term": term,
                "meaning": meaning,
                "note": note,
            }
            words.append(word)
            all_words.append(word)
        chapters.append({
            "id": f"chapter-{number}",
            "name": f"Chapter {number}",
            "title": title,
            "words": words,
        })

    book = {
        "id": "zoology-final",
        "name": "动物学词书",
        "description": "动物学期末考试核心词汇 · 按 18 章整理",
        "icon": "🦉",
        "createdAt": "2026-06-22",
        "chapters": chapters,
    }
    payload = "window.DEFAULT_BOOKS = " + json.dumps([book], ensure_ascii=False, indent=2) + ";\n"
    (ROOT / "vocabulary.js").write_text(payload, encoding="utf-8")
    print(f"Built {len(chapters)} chapters / {len(all_words)} words")


if __name__ == "__main__":
    main()
